import logging
import os
import math
import sqlite3
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application,
    CommandHandler,
    CallbackQueryHandler,
    ConversationHandler,
    MessageHandler,
    filters,
    ContextTypes,
)
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---------------- CONFIG & CONSTANTS ----------------
BOT_TOKEN = "8989501346:AAH5r7kJOEwJfBvqDAgFrXArMIcCzpCb3tg"
MAIN_DIR = "Data_Proyek"
DB_NAME = "project_data.db"

CAT_A = [
    "DD-S3-1", "DD-BSS-S1", "HB-PS-1", "DD-BTS-S1", "OS-SM-1", 
    "SLACK SUPPORT POLE", "PU-S7.0-140", "PU-S9.0-140", "PU-AS", 
    "MH-HH2", "DD-RV-C", "TC-SM-48", "SC-OF-SM-48", "PP-OF-IN"
]

ALL_DESIGNATORS = [
    "DC-OF-SM-48C", "AC-OF-SM-48C", "SC-OF-SM-48", "OS-SM-1", "TC-SM-48", 
    "PU-S7.0-140", "PU-S9.0-140", "PU-AS", "PP-OF-IN", "DD-S3-1", 
    "DD-BSS-S1", "HB-PS-1", "DD-BM-100-1", "DD-BM-HDPE-40-1", 
    "DD-HDPE-40-1", "DD-ROD", "DD-RV-1", "DD-RV-CONCRETE", "DD-RV-C",
    "MH-HH2", "BC-TR-SOIL-1", "BC-TR-SOIL-2", "BC-TR-C-1", "BC-TR-C-5", 
    "SLACK SUPPORT POLE"
]

PREDEFINED_FOLDERS = [
    "MGL - JT.01",
    "JT.01 - JT.02",
    "JT.02 - JT.03",
    "JT.03 - JT.04",
    "JT.04 - JT.05",
    "JT.05 - JT.06",
    "JT.06 - JT.07",
    "JT.07 - JT.08",
    "JT.08 - JT.09",
    "JT.09 - TEM"
]

# States
(
    START_ROUTE,
    INPUT_NEW_FOLDER,
    SELECT_DESIGNATOR,
    INPUT_VOLUME,
    UPLOAD_PHOTO,
    CEK_SELECT_FOLDER,
    CEK_SELECT_DESIGNATOR,
    CONFIRM_GENERATE,
) = range(8)


# ---------------- DATABASE LOGIC ----------------
def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS projects (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            folder_name TEXT,
            designator TEXT,
            volume REAL
        )
    ''')
    conn.commit()
    conn.close()

def save_project(folder_name, designator, volume):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('INSERT INTO projects (folder_name, designator, volume) VALUES (?, ?, ?)', 
              (folder_name, designator, volume))
    conn.commit()
    conn.close()
    
def get_folders():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT DISTINCT folder_name FROM projects ORDER BY folder_name ASC')
    rows = c.fetchall()
    conn.close()
    return [r[0] for r in rows]

def get_designators(folder_name):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT DISTINCT designator FROM projects WHERE folder_name = ? ORDER BY designator ASC', 
              (folder_name,))
    rows = c.fetchall()
    conn.close()
    return [r[0] for r in rows]


# ---------------- UTILS & VALIDATION ----------------
def get_project_dir(folder_name, designator):
    path = os.path.join(MAIN_DIR, folder_name, designator)
    os.makedirs(path, exist_ok=True)
    return path

def get_min_photos(designator, volume):
    if designator in CAT_A:
        return 1
    else:
        return math.ceil(volume / 25)

def get_sorted_photos(dir_path):
    if not os.path.exists(dir_path):
        return []
    files = [f for f in os.listdir(dir_path) if f.startswith('foto_') and f.endswith('.jpg')]
    # Sort files logically by the index number in filename
    files.sort(key=lambda x: int(x.split('_')[1].split('.')[0]))
    return [os.path.join(dir_path, f) for f in files]


# ---------------- DOCX GENERATOR ----------------
def generate_docx(folder_name, designator, image_paths, output_path):
    doc = Document()
    
    # Atur Margin Halaman agar tabel grid pas dan tidak overflow
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        
    doc.add_heading(f'Dokumentasi Proyek: {folder_name}', 0)
    doc.add_heading(f'Designator: {designator}', 1)
    
    # Aturan ketat 2x2 grid (Maks 4 foto per halaman)
    chunks = [image_paths[i:i + 4] for i in range(0, len(image_paths), 4)]
    
    for chunk_idx, chunk in enumerate(chunks):
        if chunk_idx > 0:
            doc.add_page_break()
            
        # Membuat tabel 2 baris x 2 kolom
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        table.autofit = False
        
        # Set fix width (8 cm per cell, proporsional dengan kertas A4 dan margin)
        for col in table.columns:
            col.width = Cm(8)
            
        for idx, img_path in enumerate(chunk):
            row_idx = idx // 2
            col_idx = idx % 2
            cell = table.cell(row_idx, col_idx)
            
            # Format Paragraf di dalam cell untuk center alignment
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            try:
                # Resize gambar menjadi 7.5cm (sedikit lebih kecil dari cell)
                # Tinggi akan menyesuaikan secara proporsional otomatis
                run.add_picture(img_path, width=Cm(7.5))
            except Exception as e:
                run.add_text(f"[Gagal meload gambar: {os.path.basename(img_path)}]")
                
    doc.save(output_path)


# ---------------- BOT HANDLERS ----------------
async def start_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    keyboard = [
        [
            InlineKeyboardButton("Cek Data", callback_data="cek_data"),
            InlineKeyboardButton("Input Data Baru", callback_data="input_baru"),
        ]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    msg = "Selamat datang! Silakan pilih menu berikut:"
    
    if update.message:
        await update.message.reply_text(msg, reply_markup=reply_markup)
    else:
        await update.callback_query.edit_message_text(msg, reply_markup=reply_markup)
        
    return START_ROUTE

async def handle_start_route(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query = update.callback_query
    await query.answer()
    
    if query.data == "cek_data":
        context.user_data['folder_list'] = PREDEFINED_FOLDERS
        keyboard = []
        for i in range(0, len(PREDEFINED_FOLDERS), 2):
            row = [InlineKeyboardButton(PREDEFINED_FOLDERS[j], callback_data=f"foldx_{j}") for j in range(i, min(i+2, len(PREDEFINED_FOLDERS)))]
            keyboard.append(row)
        keyboard.append([InlineKeyboardButton("Kembali", callback_data="back_start")])
        
        await query.edit_message_text("Pilih Folder Proyek:", reply_markup=InlineKeyboardMarkup(keyboard))
        return CEK_SELECT_FOLDER
        
    elif query.data == "input_baru":
        keyboard = []
        for i in range(0, len(PREDEFINED_FOLDERS), 2):
            row = [InlineKeyboardButton(PREDEFINED_FOLDERS[j], callback_data=f"newfold_{j}") for j in range(i, min(i+2, len(PREDEFINED_FOLDERS)))]
            keyboard.append(row)
        keyboard.append([InlineKeyboardButton("Kembali", callback_data="back_start")])
        await query.edit_message_text("Pilih Nama Folder Proyek:", reply_markup=InlineKeyboardMarkup(keyboard))
        return INPUT_NEW_FOLDER

    elif query.data == "back_start":
        return await start_cmd(update, context)
        
    return START_ROUTE

async def input_new_folder(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query = update.callback_query
    await query.answer()
    
    if query.data == "back_start":
        return await start_cmd(update, context)
        
    idx = int(query.data.replace("newfold_", ""))
    folder_name = PREDEFINED_FOLDERS[idx]
    context.user_data['folder_name'] = folder_name
    
    keyboard = []
    # Membelah array menjadi 2 kolom untuk inline keyboard
    for i in range(0, len(ALL_DESIGNATORS), 2):
        row = [InlineKeyboardButton(desc, callback_data=f"desig_{desc}") for desc in ALL_DESIGNATORS[i:i+2]]
        keyboard.append(row)
        
    await query.edit_message_text(
        f"Folder '{folder_name}' dipilih.\nSilakan pilih Designator:", 
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return SELECT_DESIGNATOR

async def select_designator(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query = update.callback_query
    await query.answer()
    
    designator = query.data.replace("desig_", "")
    context.user_data['designator'] = designator
    
    await query.edit_message_text(f"Designator terpilih: {designator}\n\nMasukkan Volume (berupa angka):")
    return INPUT_VOLUME

async def input_volume(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    try:
        volume = float(update.message.text)
        context.user_data['volume'] = volume
    except ValueError:
        await update.message.reply_text("Volume harus berupa angka. Silakan ketik ulang Volume:")
        return INPUT_VOLUME
        
    folder_name = context.user_data['folder_name']
    designator = context.user_data['designator']
    
    # Save project data to DB
    save_project(folder_name, designator, volume)
    
    # Ensure dir exists
    get_project_dir(folder_name, designator)
    context.user_data['uploaded_count'] = 0
    
    keyboard = [[InlineKeyboardButton("Selesai Upload", callback_data="selesai_upload")]]
    await update.message.reply_text(
        f"Data tersimpan (Volume: {volume}).\n\n"
        "Silakan upload foto satu per satu.\n"
        "Pastikan kirim foto sebagai 'Photo' (Compress), bukan sebagai Document/File.\n"
        "Jika seluruh foto sudah selesai dikirim, klik tombol 'Selesai Upload'.",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return UPLOAD_PHOTO

async def handle_photo_upload(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    folder_name = context.user_data['folder_name']
    designator = context.user_data['designator']
    
    photo_file = await update.message.photo[-1].get_file()
    
    dir_path = get_project_dir(folder_name, designator)
    existing_files = [f for f in os.listdir(dir_path) if f.startswith('foto_') and f.endswith('.jpg')]
    next_index = len(existing_files) + 1
    file_path = os.path.join(dir_path, f"foto_{next_index}.jpg")
    
    await photo_file.download_to_drive(file_path)
    context.user_data['uploaded_count'] += 1
    
    await update.message.reply_text(f"✅ Foto ke-{next_index} berhasil disimpan.")
    return UPLOAD_PHOTO

async def handle_finish_upload(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query = update.callback_query
    await query.answer()
    
    folder_name = context.user_data['folder_name']
    designator = context.user_data['designator']
    volume = context.user_data.get('volume', 0)
    
    min_photos = get_min_photos(designator, volume)
    dir_path = get_project_dir(folder_name, designator)
    total_photos = len(get_sorted_photos(dir_path))
    
    if total_photos < min_photos:
        await query.message.reply_text(
            f"❌ Validasi Gagal!\n"
            f"Designator {designator} (Volume: {volume}) wajib memiliki minimal {min_photos} foto.\n"
            f"Saat ini folder baru terisi {total_photos} foto.\n\n"
            f"Silakan upload foto tambahannya, lalu tekan kembali tombol 'Selesai Upload' di atas jika sudah lengkap."
        )
        return UPLOAD_PHOTO
        
    keyboard = [
        [InlineKeyboardButton("Ya, Generate DOCX Sekarang", callback_data="generate_yes")],
        [InlineKeyboardButton("Tidak, Simpan Saja", callback_data="generate_no")]
    ]
    await query.message.reply_text(
        f"✅ Validasi Sukses!\n(Total foto di server: {total_photos} foto. Syarat minimum: {min_photos} foto).\n\n"
        "Apakah Anda ingin langsung men-generate dokumen Word (.docx) sekarang?",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return CONFIRM_GENERATE

async def cek_select_folder(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query = update.callback_query
    await query.answer()
    
    if query.data == "back_start":
        return await start_cmd(update, context)
        
    idx = int(query.data.replace("foldx_", ""))
    folder_name = context.user_data['folder_list'][idx]
    context.user_data['cek_folder'] = folder_name
    
    base_path = os.path.join(MAIN_DIR, folder_name)
    if os.path.exists(base_path):
        existing_folders = [f for f in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, f))]
        # Filter all valid designators (CAT_A & CAT_B) that exist physically
        designators = [d for d in existing_folders if d in ALL_DESIGNATORS]
    else:
        designators = []
        
    context.user_data['desig_list'] = designators
    
    if not designators:
        keyboard = [[InlineKeyboardButton("Kembali", callback_data="back_folders")]]
        await query.edit_message_text(f"Folder: {folder_name}\nBelum ada sub-folder/data foto designator di sini.", reply_markup=InlineKeyboardMarkup(keyboard))
        return CEK_SELECT_DESIGNATOR
        
    keyboard = []
    for i in range(0, len(designators), 2):
        row = [InlineKeyboardButton(designators[j], callback_data=f"cdesx_{j}") for j in range(i, min(i+2, len(designators)))]
        keyboard.append(row)
    keyboard.append([InlineKeyboardButton("Kembali", callback_data="back_folders")])
    
    await query.edit_message_text(
        f"Folder: {folder_name}\nPilih Designator (Terdeteksi di PC):", 
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return CEK_SELECT_DESIGNATOR

async def cek_select_designator(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query = update.callback_query
    await query.answer()
    
    if query.data == "back_folders":
        folders = context.user_data['folder_list']
        keyboard = []
        for i in range(0, len(folders), 2):
            row = [InlineKeyboardButton(folders[j], callback_data=f"foldx_{j}") for j in range(i, min(i+2, len(folders)))]
            keyboard.append(row)
        keyboard.append([InlineKeyboardButton("Kembali", callback_data="back_start")])
        await query.edit_message_text("Pilih Folder Proyek:", reply_markup=InlineKeyboardMarkup(keyboard))
        return CEK_SELECT_FOLDER
        
    idx = int(query.data.replace("cdesx_", ""))
    designator = context.user_data['desig_list'][idx]
    folder_name = context.user_data['cek_folder']
    
    context.user_data['folder_name'] = folder_name
    context.user_data['designator'] = designator
    
    dir_path = get_project_dir(folder_name, designator)
    total_photos = len(get_sorted_photos(dir_path))
    
    keyboard = [
        [InlineKeyboardButton("Generate DOCX", callback_data="generate_yes")],
        [InlineKeyboardButton("Kembali ke Menu Utama", callback_data="generate_no")]
    ]
    
    await query.edit_message_text(
        f"📋 Rincian Data Proyek:\n\n"
        f"Folder: {folder_name}\n"
        f"Designator: {designator}\n"
        f"Total Foto Disimpan: {total_photos} foto\n\n"
        f"Pilih aksi selanjutnya:",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return CONFIRM_GENERATE

async def confirm_generate(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query = update.callback_query
    await query.answer()
    
    if query.data == "generate_yes":
        folder_name = context.user_data['folder_name']
        designator = context.user_data['designator']
        
        await query.edit_message_text("⏳ Sedang menyusun dokumen, mohon tunggu...")
        
        output_filename = f"Report_{folder_name}_{designator}.docx".replace(" ", "_")
        output_path = os.path.join(MAIN_DIR, folder_name, designator, output_filename)
        dir_path = get_project_dir(folder_name, designator)
        
        image_paths = get_sorted_photos(dir_path)
        
        if not image_paths:
            await query.message.reply_text("❌ Tidak ada foto di folder terkait untuk di-generate!")
            return await start_cmd(update, context)
            
        # Panggil fungsi generate tabel grid 2x2 docx
        generate_docx(folder_name, designator, image_paths, output_path)
        
        # Kirim dokumen kembali ke user
        with open(output_path, 'rb') as doc_file:
            await context.bot.send_document(chat_id=query.message.chat_id, document=doc_file)
            
        await query.message.reply_text("✅ Pembuatan dokumen sukses!")
    else:
        await query.edit_message_text("🆗 Kembali ke menu utama. Data aman tersimpan.")
        
    return await start_cmd(update, context)

def main():
    # Setup Database
    init_db()
    
    # Initialize bot
    app = Application.builder().token(BOT_TOKEN).build()
    
    # Conversation Handler Setup
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start_cmd)],
        states={
            START_ROUTE: [CallbackQueryHandler(handle_start_route)],
            INPUT_NEW_FOLDER: [CallbackQueryHandler(input_new_folder, pattern="^(newfold_|back_start)")],
            SELECT_DESIGNATOR: [CallbackQueryHandler(select_designator, pattern="^desig_")],
            INPUT_VOLUME: [MessageHandler(filters.TEXT & ~filters.COMMAND, input_volume)],
            UPLOAD_PHOTO: [
                MessageHandler(filters.PHOTO, handle_photo_upload),
                CallbackQueryHandler(handle_finish_upload, pattern="^selesai_upload$")
            ],
            CEK_SELECT_FOLDER: [CallbackQueryHandler(cek_select_folder, pattern="^(foldx_|back_start)")],
            CEK_SELECT_DESIGNATOR: [CallbackQueryHandler(cek_select_designator, pattern="^(cdesx_|back_folders)")],
            CONFIRM_GENERATE: [CallbackQueryHandler(confirm_generate, pattern="^generate_")],
        },
        fallbacks=[CommandHandler("start", start_cmd)],
    )
    
    app.add_handler(conv_handler)
    
    print("Bot is running... Silakan cek Telegram Anda.")
    app.run_polling()

if __name__ == "__main__":
    main()
