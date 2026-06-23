from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io

def generate_report_docx(
    report_title: str,
    project_data: dict,
    image_streams: list, # List of io.BytesIO containing downloaded images
) -> io.BytesIO:
    """
    Membuat dokumen Word (.docx) sesuai dengan format Foto 1.
    """
    doc = Document()
    
    # 1. Judul Dokumen (Tengah, Bold, Underline)
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(f'"{report_title.upper()}"')
    run.bold = True
    run.underline = True
    run.font.size = Pt(16)
    
    doc.add_paragraph() # Spasi
    
    # 2. Tabel Informasi Detail (Borderless)
    # Membuat tabel 6 baris x 3 kolom (Label, Titik Dua, Value)
    table = doc.add_table(rows=6, cols=3)
    table.autofit = False
    
    # Atur lebar kolom (misal: 4cm, 0.5cm, 10cm)
    col_widths = [Cm(4), Cm(0.5), Cm(10)]
    for i in range(3):
        for cell in table.columns[i].cells:
            cell.width = col_widths[i]

    details = [
        ("Proyek", project_data.get('proyek', '-')),
        ("No Kontrak", project_data.get('no_kontrak', '-')),
        ("Nomor PO", project_data.get('no_po', '-')),
        ("Lokasi", project_data.get('area', '-')),
        ("Site Operation", project_data.get('site_operation', '-')),
        ("Pelaksana", project_data.get('pelaksana', '-'))
    ]
    
    for i, (label, val) in enumerate(details):
        row_cells = table.rows[i].cells
        
        # Label (Bold)
        r1 = row_cells[0].paragraphs[0].add_run(label)
        r1.bold = True
        
        # Titik dua (Bold)
        r2 = row_cells[1].paragraphs[0].add_run(":")
        r2.bold = True
        
        # Value (Bold)
        r3 = row_cells[2].paragraphs[0].add_run(str(val))
        r3.bold = True

    doc.add_paragraph() # Spasi sebelum foto
    
    # Menambahkan garis bawah tabel untuk pemisah (Opsional, tapi di foto terlihat ada border bawah)
    # Untuk implementasi sederhana, kita lewati border bawah manual agar tidak terlalu rumit di python-docx
    
    # 3. Grid Foto (2 kolom bersebelahan)
    # Kita buat tabel tak terlihat untuk menyusun foto 2 kolom
    num_images = len(image_streams)
    rows_needed = (num_images + 1) // 2
    
    img_table = doc.add_table(rows=rows_needed, cols=2)
    img_table.autofit = True
    
    img_idx = 0
    for r in range(rows_needed):
        for c in range(2):
            if img_idx < num_images:
                cell = img_table.cell(r, c)
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                
                # Tambahkan gambar (lebar disesuaikan agar pas 2 kolom, misal 3 inch)
                try:
                    run.add_picture(image_streams[img_idx], width=Inches(3.0))
                except Exception as e:
                    print(f"Error adding picture {img_idx}: {e}")
                    
                img_idx += 1
    
    # Simpan ke stream (tidak ke file fisik, agar bisa langsung dikirim bot)
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream
